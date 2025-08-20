import streamlit as st
import analytics

analytics.load_analytics()
st.set_page_config(
    page_title="Notexa -Learning Made Simple.",
    page_icon="assets/logo.png"
)
import requests
import os
import tempfile
import fitz  # PyMuPDF
from docx import Document
from io import BytesIO

# --- Gemini API Setup ---
API_KEY = os.getenv("API_KEY")

def ask_gemini(text):
    prompt = (
        "Generate academic quiz from this academic notes, Quiz should have 3 sections, section A(must include 8 multiple choice questions), section B(must include 4 questions, 2 of True or False and 2 questions of Fill in blanks) and Section C(must include 2 open-ended questions):\n\n"
        f"{text}\n\nQuiz:"
    )
    url = f"https://generativelanguage.googleapis.com/v1/models/gemini-2.5-pro:generateContent?key=AIzaSyAY0FuRkzB9yzsdi9v40g7OjPME_f2WWik"
    headers = {"Content-Type": "application/json"}
    data = {"contents": [{"parts": [{"text": prompt}]}]}
    try:
        res = requests.post(url, headers=headers, json=data)
        res.raise_for_status()
        return res.json()['candidates'][0]['content']['parts'][0]['text']
    except Exception as e:
        return f"❌ Gemini API error: {e}"

def extract_text(file):
    if file.type == "application/pdf":
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                tmp.write(file.read())
                tmp_path = tmp.name
            doc = fitz.open(tmp_path)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            os.remove(tmp_path)
            return text.strip()
        except Exception:
            return None
    elif file.type == "text/plain":
        try:
            return file.read().decode("utf-8").strip()
        except Exception:
            return None
    elif file.type in [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword"
    ] or file.name.endswith(".docx"):
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                tmp.write(file.read())
                tmp_path = tmp.name
            doc = Document(tmp_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            os.remove(tmp_path)
            return text.strip()
        except Exception:
            return None
    else:
        return None

def summary_to_docx(summary):
    doc = Document()
    for line in summary.splitlines():
        if line.strip().startswith("# "):
            doc.add_heading(line.replace("#", "").strip(), level=1)
        elif line.strip().startswith("## "):
            doc.add_heading(line.replace("#", "").strip(), level=2)
        elif line.strip().startswith("### "):
            doc.add_heading(line.replace("#", "").strip(), level=3)
        elif line.strip().startswith("* ") or line.strip().startswith("- "):
            doc.add_paragraph(line.strip()[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- UI ---
st.title("📚 🧪 Quiz Generator from Notes")
st.caption("Turn your notes into practice quizzes with AI")

st.markdown(
    "Upload your academic notes (PDF, TXT, or DOCX). "
    "Get a practice quiz generated from your notes with AI. "
)

st.markdown("")

uploaded_file = st.file_uploader(
    "Choose a notes file (PDF, TXT, DOCX)",
    type=["pdf", "txt", "docx"],
    help="Upload your notes file here."
)

summary = ""
if uploaded_file:
    text = extract_text(uploaded_file)
    if not text:
        st.warning("Sorry, we couldn't read your file. Please upload a valid PDF, TXT, or DOCX file.")
    else:
        if st.button("Generate Quiz", use_container_width=True):
            with st.spinner("Generating Quiz..."):
                summary = ask_gemini(text)
                if not summary or summary.startswith("❌"):
                    st.warning("Sorry, we couldn't generate a Quiz. Please try again.")
                else:
                    st.success("Quiz ready!")

if summary:
    st.markdown("### 📝 Quiz")
    st.markdown(summary, unsafe_allow_html=True)
    col1, col2 = st.columns(2)
    with col1:
        # (Copy button code as before, or use JS solution for clipboard)
        pass
    with col2:
        docx_file = summary_to_docx(summary)
        st.download_button(
            "Download Quiz",
            docx_file,
            file_name="summary.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

# --- Go Premium Button in Sidebar ---
st.sidebar.markdown('''<hr style="margin-top:2em;margin-bottom:0.5em;">''', unsafe_allow_html=True)
st.sidebar.markdown('''
    <a href="/premium" target="_self" style="
        display: block;
        background: linear-gradient(90deg, #FFD700 0%, #FFB300 100%);
        color: #222;
        padding: 0.5em 1.2em;
        border-radius: 2em;
        font-weight: bold;
        font-size: 1.1em;
        text-align: center;
        text-decoration: none;
        box-shadow: 0 2px 8px rgba(0,0,0,0.08);
        border: 2px solid #fff2b2;
        margin-top: 0.5em;
        margin-bottom: 1em;
        transition: background 0.2s;
    " onmouseover="this.style.background='#FFB300'" onmouseout="this.style.background='linear-gradient(90deg, #FFD700 0%, #FFB300 100%)'">
        🙏 Donate
    </a>
''', unsafe_allow_html=True)
st.sidebar.markdown(
    '<div style="margin-top:3em; margin-bottom:1em; font-size:1.05em;">'
    '📢 <b>Turn students and teachers into customers.</b> '
    '<a href="https://your-ad-link.com" target="_blank" style="color:#00b894; font-weight:bold; text-decoration:underline;">Advertise on Notexa</a>.'
    '</div>',
    unsafe_allow_html=True
)
st.markdown('''
    <style>
    .notexa-footer {
        position: fixed;
        left: 0;
        bottom: 0;
        width: 100vw;
        background: rgba(255,255,255,0.0);
        color: #888;
        text-align: center;
        font-size: 1em;
        padding: 0.7em 0 0.5em 0;
        z-index: 9999;
    }
    </style>
    <div class="notexa-footer">© 2025. Made with ❤️ by the Hirwa Leon</div>
    ''', unsafe_allow_html=True)
