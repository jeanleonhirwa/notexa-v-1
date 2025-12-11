import streamlit as st
import analytics

analytics.load_analytics()
st.set_page_config(
    page_title="Notexa -Learning Made Simple.",
    page_icon="assets/logo.png"
)
import os
import tempfile
import fitz  # PyMuPDF
from docx import Document
from io import BytesIO
import sys

# Add parent directory to path to import gemini_api
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from gemini_api import summarize_notes

def extract_text(file):
    if file.type == "application/pdf":
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                tmp.write(file.read())
                tmp_path = tmp.name
            doc = fitz.open(tmp_path)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            os.remove(tmp_path)
            return text.strip()
        except Exception:
            return None
    elif file.type == "text/plain":
        try:
            return file.read().decode("utf-8").strip()
        except Exception:
            return None
    elif file.type in [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword"
    ] or file.name.endswith(".docx"):
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                tmp.write(file.read())
                tmp_path = tmp.name
            doc = Document(tmp_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            os.remove(tmp_path)
            return text.strip()
        except Exception:
            return None
    else:
        return None

def summary_to_docx(summary):
    doc = Document()
    for line in summary.splitlines():
        if line.strip().startswith("# "):
            doc.add_heading(line.replace("#", "").strip(), level=1)
        elif line.strip().startswith("## "):
            doc.add_heading(line.replace("#", "").strip(), level=2)
        elif line.strip().startswith("### "):
            doc.add_heading(line.replace("#", "").strip(), level=3)
        elif line.strip().startswith("* ") or line.strip().startswith("- "):
            doc.add_paragraph(line.strip()[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- UI ---
st.title("📚 AI Note Summarizer")
st.caption("Summarize your notes into easy-to-read points.")

st.markdown(
    "Upload your academic notes (PDF, TXT, or DOCX). "
    "Get a simple summary anyone can understand!"
)

st.markdown("")

uploaded_file = st.file_uploader(
    "Choose a notes file (PDF, TXT, DOCX)",
    type=["pdf", "txt", "docx"],
    help="Upload your notes file here."
)

summary = ""
if uploaded_file:
    text = extract_text(uploaded_file)
    if not text:
        st.warning("Sorry, we couldn't read your file. Please upload a valid PDF, TXT, or DOCX file.")
    else:
        if st.button("Summarize Notes", use_container_width=True):
            with st.spinner("Summarizing your notes..."):
                summary = summarize_notes(text)
                if not summary or summary.startswith("❌"):
                    st.warning("Sorry, we couldn't generate a summary. Please try again.")
                else:
                    st.success("Summary ready!")

if summary:
    st.markdown("### 📝 Summary")
    st.markdown(summary, unsafe_allow_html=True)
    col1, col2 = st.columns(2)
    with col1:
        # (Copy button code as before, or use JS solution for clipboard)
        pass
    with col2:
        docx_file = summary_to_docx(summary)
        st.download_button(
            "Download Summary",
            docx_file,
            file_name="summary.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

# --- Go Premium Button in Sidebar ---
st.sidebar.markdown('''<hr style="margin-top:2em;margin-bottom:0.5em;">''', unsafe_allow_html=True)
st.sidebar.markdown('''
    <a href="/premium" target="_self" style="
        display: block;
        background: linear-gradient(90deg, #FFD700 0%, #FFB300 100%);
        color: #222;
        padding: 0.5em 1.2em;
        border-radius: 2em;
        font-weight: bold;
        font-size: 1.1em;
        text-align: center;
        text-decoration: none;
        box-shadow: 0 2px 8px rgba(0,0,0,0.08);
        border: 2px solid #fff2b2;
        margin-top: 0.5em;
        margin-bottom: 1em;
        transition: background 0.2s;
    " onmouseover="this.style.background='#FFB300'" onmouseout="this.style.background='linear-gradient(90deg, #FFD700 0%, #FFB300 100%)'">
        🙏 Donate
    </a>
''', unsafe_allow_html=True)
st.sidebar.markdown(
    '<div style="margin-top:3em; margin-bottom:1em; font-size:1.05em;">'
    '📢 <b>Turn students and teachers into customers.</b> '
    '<a href="https://your-ad-link.com" target="_blank" style="color:#00b894; font-weight:bold; text-decoration:underline;">Advertise on Notexa</a>.'
    '</div>',
    unsafe_allow_html=True
)
st.markdown('''
    <style>
    .notexa-footer {
        position: fixed;
        left: 0;
        bottom: 0;
        width: 100vw;
        background: rgba(255,255,255,0.0);
        color: #888;
        text-align: center;
        font-size: 1em;
        padding: 0.7em 0 0.5em 0;
        z-index: 9999;
    }
    </style>
    <div class="notexa-footer">© 2025. Made with ❤️ by the Hirwa Leon</div>
    ''', unsafe_allow_html=True)
