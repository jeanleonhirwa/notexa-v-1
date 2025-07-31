import streamlit as st
import requests
import os
import tempfile
import fitz  # PyMuPDF
from docx import Document
from io import BytesIO

# --- Gemini API Setup ---
API_KEY = os.getenv("API_KEY")

def ask_gemini(text):
    prompt = (
        "Summarize the following all academic notes in simple, clear English that is easy to understand:\n\n"
        f"{text}\n\nSummary:"
    )
    url = f"https://generativelanguage.googleapis.com/v1/models/gemini-2.5-pro:generateContent?key={API_KEY}"
    headers = {"Content-Type": "application/json"}
    data = {"contents": [{"parts": [{"text": prompt}]}]}
    try:
        res = requests.post(url, headers=headers, json=data)
        res.raise_for_status()
        return res.json()['candidates'][0]['content']['parts'][0]['text']
    except Exception as e:
        return f"❌ Gemini API error: {e}"

def extract_text(file):
    if file.type == "application/pdf":
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                tmp.write(file.read())
                tmp_path = tmp.name
            doc = fitz.open(tmp_path)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            os.remove(tmp_path)
            return text.strip()
        except Exception:
            return None
    elif file.type == "text/plain":
        try:
            return file.read().decode("utf-8").strip()
        except Exception:
            return None
    elif file.type in [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword"
    ] or file.name.endswith(".docx"):
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                tmp.write(file.read())
                tmp_path = tmp.name
            doc = Document(tmp_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            os.remove(tmp_path)
            return text.strip()
        except Exception:
            return None
    else:
        return None

def summary_to_docx(summary):
    doc = Document()
    for line in summary.splitlines():
        if line.strip().startswith("# "):
            doc.add_heading(line.replace("#", "").strip(), level=1)
        elif line.strip().startswith("## "):
            doc.add_heading(line.replace("#", "").strip(), level=2)
        elif line.strip().startswith("### "):
            doc.add_heading(line.replace("#", "").strip(), level=3)
        elif line.strip().startswith("* ") or line.strip().startswith("- "):
            doc.add_paragraph(line.strip()[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- UI ---
st.title("📚 AI Note Summarizer")
st.caption("Summarize your notes into easy-to-read points.")

st.markdown(
    "Upload your academic notes (PDF, TXT, or DOCX). "
    "Get a simple summary anyone can understand!"
)

st.markdown("")

uploaded_file = st.file_uploader(
    "Choose a notes file (PDF, TXT, DOCX)",
    type=["pdf", "txt", "docx"],
    help="Upload your notes file here."
)

summary = ""
if uploaded_file:
    text = extract_text(uploaded_file)
    if not text:
        st.warning("Sorry, we couldn't read your file. Please upload a valid PDF, TXT, or DOCX file.")
    else:
        if st.button("Summarize Notes", use_container_width=True):
            with st.spinner("Summarizing your notes..."):
                summary = ask_gemini(text)
                if not summary or summary.startswith("❌"):
                    st.warning("Sorry, we couldn't generate a summary. Please try again.")
                else:
                    st.success("Summary ready!")

if summary:
    st.markdown("### 📝 Summary")
    st.markdown(summary, unsafe_allow_html=True)
    col1, col2 = st.columns(2)
    with col1:
        # (Copy button code as before, or use JS solution for clipboard)
        pass
    with col2:
        docx_file = summary_to_docx(summary)
        st.download_button(
            "Download Summary",
            docx_file,
            file_name="summary.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )